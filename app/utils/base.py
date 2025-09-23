import json
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, Any

from app.config import settings
from app.utils.console import confirm_prompt


def is_supported(file: Path) -> bool:
	"""Проверить, что файл имеет поддерживаемое расширение"""
	return file.suffix.lower() in settings.SUPPORTED_FORMATS


def extract_data_from_file(file_path: Path) -> tuple[str, Optional[list[dict[str, Any]]]]:
	"""
	Извлекает текст и таблицы из файла. Автоматически выбирает метод по расширению.

	Args:
		file_path: Путь к файлу

	Returns:
		Кортеж (текст, список таблиц)
		Таблицы: список словарей с данными таблиц
	"""
	if not file_path.exists():
		raise FileNotFoundError(f"Файл не найден: {file_path}")

	file_ext = file_path.suffix.lower()

	try:
		if file_ext == '.pdf':
			return _extract_from_pdf(file_path)
		elif file_ext in ['.docx', '.doc']:
			return _extract_from_docx(file_path)
		elif file_ext == '.txt':
			return _extract_from_txt(file_path)
		else:
			raise ValueError(f"Неподдерживаемый формат файла: {file_ext}")
	except Exception as e:
		raise Exception(f"Ошибка извлечения данных из {file_path.name}: {e}")


def _extract_from_pdf(file_path: Path) -> tuple[str, Optional[list[dict[str, Any]]]]:
	"""Извлекает текст и таблицы из PDF"""
	import pdfplumber

	text_content = ""
	tables_data = []

	with pdfplumber.open(file_path) as pdf:
		for page in pdf.pages:
			# Извлекаем текст
			page_text = page.extract_text()
			if page_text:
				text_content += page_text + "\n"

			# Извлекаем таблицы
			page_tables = page.extract_tables()
			for i, table in enumerate(page_tables):
				if table and any(any(cell for cell in row) for row in table):
					tables_data.append({
						'page': page.page_number,
						'table_number': i + 1,
						'data': table,
						'source': file_path.name
					})

	return text_content, tables_data if tables_data else None


def _extract_from_docx(file_path: Path) -> tuple[str, Optional[list[dict[str, Any]]]]:
	"""Извлекает текст и таблицы из DOCX"""
	from docx import Document

	doc = Document(str(file_path))
	text_content = ""
	tables_data = []

	# Извлекаем текст из параграфов
	for paragraph in doc.paragraphs:
		if paragraph.text.strip():
			text_content += paragraph.text + "\n"

	# Извлекаем таблицы
	for i, table in enumerate(doc.tables):
		table_data = []
		for row in table.rows:
			row_data = [cell.text.strip() for cell in row.cells]
			if any(row_data):
				table_data.append(row_data)

		if table_data:
			tables_data.append({
				'table_number': i + 1,
				'data': table_data,
				'source': file_path.name
			})

	return text_content, tables_data if tables_data else None


def _extract_from_txt(file_path: Path) -> tuple[str, None]:
	"""Извлекает текст из TXT (таблиц нет)"""
	with open(file_path, 'r', encoding='utf-8') as f:
		text_content = f.read()
	return text_content, None


def get_current_year() -> int:
	"""
	Возвращает текущий год.
	Используется как значение по умолчанию при парсинге.
	"""
	return datetime.now().year


def parse_range_string(range_str: str | None, total: int) -> tuple[int, int]:
	"""
	Парсит строку диапазона в offset и limit.

	:param range_str: Строка диапазона ('1-10', ':10', '5:', etc.)
	:param total: Общее количество документов
	:return: Кортеж (offset, limit)
	"""
	if not range_str or range_str.strip().lower() == "all":
		return 0, total  # Все документы

	range_str = range_str.strip()

	# Приводим к нормальному виду: убираем лишние пробелы и приводим знаки '-' и ':' к одному стилю
	normalized_range = range_str.replace(':', '-').replace(' ', '')

	# Регулярное выражение для извлечения частей диапазона
	match = re.match(r'^(\d*)\s*-\s*(\d*)$', normalized_range)
	if match:
		start_part, end_part = match.groups()

		try:
			# Начало диапазона (1-based → 0-based)
			start = int(start_part) - 1 if start_part else 0
			# Конец диапазона
			end = int(end_part) if end_part else total

			# Проверка корректности
			if start < 0:
				raise ValueError(f"Начало диапазона не может быть отрицательным: {start + 1}")
			if end_part and end <= start:
				raise ValueError(f"Конец диапазона должен быть больше начала: {start + 1}-{end}")
			if start >= total:
				raise ValueError(f"Начало диапазона ({start + 1}) превышает общее количество документов ({total})")

			# Ограничение (разница между началом и концом)
			limit = end - start if end > start else None

		except ValueError as e:
			raise ValueError(f"Некорректный диапазон '{range_str}': {e}")

	elif normalized_range.isdigit():  # Просто число — ограниченное количество документов
		limit = int(normalized_range)
		if limit <= 0:
			raise ValueError(f"Лимит должен быть положительным: {limit}")
		return 0, limit

	else:
		raise ValueError(f"Некорректный формат диапазона: '{range_str}'. Используйте: '1-10', ':10', '5-', 'all'")

	return start, limit


def get_unique_filename(
		directory: Path,
		base_name: str,
		postfix: str = "",
		extension: str = ".xlsx",
		force_overwrite: bool = False
) -> Path:
	"""
	Генерирует уникальное имя файла, добавляя индекс если файл уже существует.

	Args:
		directory: Папка для сохранения
		base_name: Базовое имя файла (без расширения)
		extension: Расширение файла (по умолчанию .xlsx)
		postfix: Дополнительное окончание к имени файла
		force_overwrite: Принудительная перезапись существующего файла

	Returns:
		Путь к уникальному файлу
	"""
	# Если force_overwrite=True, просто возвращаем путь без проверок
	if force_overwrite:
		filename = f"{base_name}{postfix}{extension}"
		return directory / filename

	counter = 1
	while True:
		if counter == 1:
			filename = f"{base_name}{postfix}{extension}"
		else:
			filename = f"{base_name}{postfix}-{counter:02d}{extension}"

		file_path = directory / filename

		if not file_path.exists():
			return file_path

		# Предлагаем пользователю выбрать действие
		if confirm_prompt(
				f"Файл {filename} уже существует. Перезаписать?",
				default=False  # По умолчанию "Нет" для безопасности
		):
			return file_path
		else:
			counter += 1


def format_string_list(
		string_list: list[str] | str | None,
		default_text: str = "—",
		max_line_length: Optional[int] = None,
		separator: str = "\n"
) -> str:
	"""
	Форматирует список строк для отображения.

	Args:
		string_list: Список строк для форматирования
		default_text: Текст если список пустой
		max_line_length: Максимальная длина каждой строки
		separator: Разделитель между строками

	Returns:
		Отформатированная строка
	"""
	if not string_list:
		return default_text

	# Нормализуем входные данные к списку строк
	if isinstance(string_list, str):
		try:
			# Пробуем распарсить JSON строку
			parsed = json.loads(string_list)
			lines = parsed if isinstance(parsed, list) else [parsed]
		except (json.JSONDecodeError, TypeError):
			lines = [string_list]
	else:
		lines = string_list

	if not lines:
		return default_text

	# Обрабатываем обрезание строк
	if max_line_length:
		formatted_lines = []
		for line in lines:
			line_str = str(line)  # На случай если не строка
			if len(line_str) > max_line_length:
				formatted_lines.append(line_str[:max_line_length] + "...")
			else:
				formatted_lines.append(line_str)
		return separator.join(formatted_lines)
	else:
		return separator.join(str(line) for line in lines)
