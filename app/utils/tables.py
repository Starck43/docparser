import re
from pathlib import Path
from typing import Optional

import pdfplumber

from app.utils.base import normalize_date


def _process_table(table: list[list[str]], year: int) -> list[list[str]]:
	"""
	Из исходной таблицы делает итоговый список отфильтрованный по году [дата, значения по товарам, итог].
	"""
	results = []
	header_idx = 0

	if not table or len(table) < 2:
		return []

	for idx, row in enumerate(table[1:], 1):
		if not row or not row[0]:
			continue

		date = normalize_date(row[0])
		if not date:
			continue

		if not header_idx:
			header_idx = idx

		try:
			if int(date.split(".")[1]) != year:
				continue
		except (IndexError, ValueError):
			continue

		clean_row = [cell.strip() if cell else "" for cell in row]

		# Для одной колонки НЕ добавляем итог
		if len(clean_row) <= 2:
			results.append([date] + clean_row[1:])
			continue

		numeric_values = []
		for cell in clean_row[1:]:
			if cell:
				try:
					numeric_values.append(float(cell.replace(",", ".").replace(" ", "")))
				except ValueError:
					numeric_values.append(0.0)
			else:
				numeric_values.append(0.0)

		# Проверяем, является ли последняя колонка итогом
		is_last_column_total = False
		if len(numeric_values) > 1:
			sum_without_last = sum(numeric_values[:-1])
			last_value = numeric_values[-1]
			if int(sum_without_last) == int(last_value):
				is_last_column_total = True

		if is_last_column_total:
			results.append([date] + clean_row[1:])
		else:
			results.append([date] + clean_row[1:] + [str(round(sum(numeric_values), 3))])

	if not results:
		return []

	# Замена шапки
	header = ["Срок поставки (мес/год)"]
	if header_idx > 0 and header_idx - 1 < len(table):
		header.extend(table[header_idx - 1][1:])

	# Добавляем "Итого" только если добавили итоговую колонку
	if results and len(results[0]) > len(header):
		header.append("Итого")
	elif len(header) > 1:
		header[-1] = "Итого"

	results.insert(0, header)
	return results


def _extract_text_block(text: str, start: Optional[str], end: Optional[str], include_markers: bool = True) -> str:
	"""Извлекает блок текста между маркерами (включая сами маркеры)."""
	start_pos = 0
	end_pos = len(text)

	if start:
		if match := re.search(re.escape(start), text):
			start_pos = match.start() if include_markers else match.end()

	if end:
		if match := re.search(re.escape(end), text):
			end_pos = match.end() if include_markers else match.start()

	return text[start_pos:end_pos].strip()


def _text_to_table_data(text: str) -> list[list[str]]:
	"""Преобразует текстовый блок таблицы в структурированные данные."""
	if not text:
		return []

	table_data = []
	for line in text.split('\n'):
		line = line.strip()
		if not line:
			continue

		# Ищем дату в начале строки
		date_match = re.match(r'^([а-я]+\.?\s*\d{2,4}\.?\s*г?\.?\s*)', line, re.IGNORECASE)

		if date_match:
			date_part = date_match.group(1).strip()
			rest = line[len(date_match.group(0)):].strip()
			# Более надежное разделение: 2+ пробела ИЛИ пробел между цифрами
			columns = re.split(r'\s{2,}|(?<=\d)\s+(?=\d)|(?<=[а-я])\s+(?=\d)', rest)
			table_data.append([date_part] + columns)
		else:
			# Для заголовков и строк без дат
			columns = re.split(r'\s{2,}', line)
			if columns:
				table_data.append(columns)

	return table_data


def extract_from_pdf(file_path: Path, year: int) -> tuple[str, Optional[list[list[list[str]]]]]:
	"""
	Извлекает текст и таблицы из PDF.
	"""

	# Собираем весь текст
	with pdfplumber.open(file_path) as pdf:
		full_text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())

	if not full_text:
		return "", None

	# Поиск блоков
	table_block = _extract_text_block(full_text, '2.', '3.')
	content_before = _extract_text_block(full_text, None, '2.')
	content_after = _extract_text_block(full_text, '4.', '5.')

	content_text = (content_before + "\n" + content_after).strip()

	# Обработка таблиц
	tables = []
	if table_block:
		table_data = _text_to_table_data(table_block)
		if processed := _process_table(table_data, year):
			tables.append(processed)

	return content_text, tables or None


def extract_from_docx(file_path: Path, year: int) -> tuple[str, Optional[list[list[list[str]]]]]:
	"""Извлекает текст и таблицы из DOCX"""
	from docx import Document

	doc = Document(str(file_path))
	text_content = ""
	tables = []

	for paragraph in doc.paragraphs:
		if paragraph.text.strip():
			text_content += paragraph.text + "\n"

	for table in doc.tables:
		table_data = []

		# переводим данные из объекта в список
		for row in table.rows:
			row_data = [cell.text.strip() for cell in row.cells]
			if any(row_data):  # Пропускаем полностью пустые строки
				table_data.append(row_data)

		if table_data:
			# Обрабатываем каждую таблицу отдельно
			processed_table = _process_table(table_data, year)
			if processed_table:
				tables.append(processed_table)

	return text_content, tables or None


def extract_from_txt(file_path: Path, year: int) -> tuple[str, Optional[list[list[list[str]]]]]:
	"""Извлекает текст и таблицы из TXT (таблицы изначально предполагается что разделены ';')"""
	with open(file_path, "r", encoding="utf-8") as f:
		text_content = f.read()

	lines = text_content.splitlines()
	table_data = []
	for line in lines:
		if ";" in line:
			row = [col.strip() for col in line.split(";")]
			table_data.append(row)

	results = _process_table(table_data, year) if table_data else None
	return text_content, results
